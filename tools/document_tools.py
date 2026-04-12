"""
Document Tools - анализ документов
Читает PDF, Word, Excel и отвечает на вопросы
"""
import os
from pathlib import Path


class DocumentTools:

    def read_pdf(self, path: str) -> str:
        """Read and extract text from PDF"""
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            text = ""
            for i, page in enumerate(reader.pages):
                text += f"\n--- Page {i+1} ---\n"
                text += page.extract_text() or ""
                if len(text) > 8000:
                    text += "\n... (document truncated, too long)"
                    break
            return text.strip() or "Could not extract text from PDF"
        except ImportError:
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        text += f"\n--- Page {i+1} ---\n"
                        text += page.extract_text() or ""
                        if len(text) > 8000:
                            break
                return text.strip() or "Could not extract text"
            except ImportError:
                return "Install pypdf: pip install pypdf"
        except Exception as e:
            return f"PDF error: {e}"

    def read_word(self, path: str) -> str:
        """Read Word document (.docx)"""
        try:
            import docx
            doc = docx.Document(path)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Also read tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        text += " | ".join(cells) + "\n"
            return text.strip()[:8000] or "Empty document"
        except ImportError:
            return "Install python-docx: pip install python-docx"
        except Exception as e:
            return f"Word error: {e}"

    def analyze_document(self, path: str) -> str:
        """Auto-detect and read any document"""
        path = path.strip()
        if not Path(path).exists():
            return f"File not found: {path}"
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            return self.read_pdf(path)
        elif ext in [".docx", ".doc"]:
            return self.read_word(path)
        elif ext in [".xlsx", ".xls"]:
            from tools.excel_tools import ExcelTools
            return ExcelTools().read_excel(path)
        elif ext == ".csv":
            from tools.excel_tools import ExcelTools
            return ExcelTools().read_csv(path)
        elif ext in [".txt", ".md", ".py", ".js", ".html", ".json"]:
            try:
                return Path(path).read_text(encoding="utf-8", errors="replace")[:8000]
            except Exception as e:
                return f"Read error: {e}"
        else:
            return f"Unsupported format: {ext}. Supported: PDF, DOCX, XLSX, CSV, TXT"

    def get_document_info(self, path: str) -> str:
        """Get basic info about document"""
        try:
            p = Path(path)
            if not p.exists():
                return f"File not found: {path}"
            size = p.stat().st_size
            ext = p.suffix.lower()
            size_str = f"{size//1024}KB" if size > 1024 else f"{size}B"
            info = f"📄 {p.name}\n📁 Size: {size_str}\n📋 Type: {ext}"
            if ext == ".pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(path)
                    info += f"\n📖 Pages: {len(reader.pages)}"
                except Exception:
                    pass
            return info
        except Exception as e:
            return f"Error: {e}"

    def summarize_document(self, path: str) -> str:
        """Read document and prepare for AI summarization"""
        content = self.analyze_document(path)
        if content.startswith("Error") or content.startswith("File not found") or content.startswith("Install") or content.startswith("Unsupported"):
            return content
        # Return content for AI to summarize
        name = Path(path).name
        return f"Document '{name}' content:\n\n{content}"
